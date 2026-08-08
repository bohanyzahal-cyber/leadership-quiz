# -*- coding: utf-8 -*-
"""ייבוא חוברת החומר הפתוח בחזרה מקובץ Word אל content.js.

הכיוון ההפוך ל-to_docx.py: המשתמש עורך את הוורד, וזה מחזיר את העריכות
למקור שממנו נבנה ה-HTML. כך הוורד וה-HTML נשארים אותו תוכן.

זיהוי המבנה נשען על מה ש-to_docx.py כתב:
  · שם מקטע  — 12.5pt (או שם מוכר מראש)
  · כותרת    — 10.5pt שאינה מתחילה ב-▪
  · פריט     — מתחיל ב-"▪ ", ומופרד "מונח — הגדרה"
  · טבלה     — טבלת Word
הדגשות אמיתיות (bold) מומרות חזרה ל-**...**, כדי שלא ילכו לאיבוד.

הרצה:  python from_docx.py "<נתיב לקובץ הוורד>"
"""
import io, os, re, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
import base64
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
MAXW = 680          # רוחב טור בחוברת ~85 מ"מ; מעבר לזה זה בזבוז בייטים בלבד

def encode_image(blob):
    """מכווץ לרוחב הדפסה ומחזיר data-URI. דיאגרמות נשארות PNG (קווים חדים),
       וצילומים עוברים ל-JPEG כשה-PNG יוצא כבד."""
    im = Image.open(io.BytesIO(blob))
    if im.width > MAXW:
        im = im.resize((MAXW, max(1, round(im.height * MAXW / im.width))), Image.LANCZOS)
    # שקיפות חייבת להיות מורכבת על לבן ולא מומרת ישירות ל-RGB: המרה ישירה
    # ממלאת את השקוף בשחור, וכל הדיאגרמות עם הרקע השקוף יצאו מלבנים שחורים
    # שרואים בהם רק את הקווים הלבנים.
    if im.mode in ("RGBA", "LA", "P") or "transparency" in im.info:
        im = im.convert("RGBA")
        bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
        im = Image.alpha_composite(bg, im).convert("RGB")
    elif im.mode != "RGB":
        im = im.convert("RGB")
    buf = io.BytesIO(); im.save(buf, "PNG", optimize=True)
    data, mime = buf.getvalue(), "image/png"
    if len(data) > 55000:
        b2 = io.BytesIO(); im.save(b2, "JPEG", quality=78, optimize=True)
        if len(b2.getvalue()) < len(data):
            data, mime = b2.getvalue(), "image/jpeg"
    return "data:%s;base64,%s" % (mime, base64.b64encode(data).decode()), im.width, im.height

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = sys.argv[1] if len(sys.argv) > 1 else None
OUT = os.path.join(HERE, "content.js")

# שם מקטע -> (צבע, תת-כותרת). הסדר כאן הוא סדר המקטעים בחוברת.
SECTIONS = [
 ("יסודות וקלאסיות",     "#37474f", "התנהגות ארגונית · פיול · טיילור · בירוקרטיה"),
 ("אורגניות ומודרניות",  "#2e7d32", "האות'ורן · מאיו · X/Y · TQM · מערכת · תלות"),
 ("אדיג'ס",              "#6a1b9a", "PAEI · ארבעה סגנונות · מחזור חיים"),
 ("מנהיגות ותכונות",     "#ad1457", "הגדרות · כוח · התכונות · אוהיו · הסריג"),
 ("מצבית וטווח מלא",     "#ef6c00", "פידלר · הרסי ובלנשארד · בס ואבוליו"),
 ("צוות, אישיות ושאלות", "#4527a0", "טאקמן · קבוצה וצוות · Big Five · שאלות"),
]
SEC_BY_NAME = {n: (c, s) for n, c, s in SECTIONS}
# הכותרת שהמשתמש נתן למקטע הרביעי, אחרי שערך אותו ידנית
ALIAS = {"תיאוריית התכונות - מנהיגות ותכונות": "מנהיגות ותכונות"}

def rich(par):
    """מרכיב מחדש את הטקסט, כשריצות מודגשות עוטפות ב-**."""
    out = []
    for r in par.runs:
        t = r.text
        if not t: continue
        if r.bold and t.strip():
            lead = len(t) - len(t.lstrip()); trail = len(t) - len(t.rstrip())
            out.append(t[:lead] + "**" + t.strip() + "**" + (t[len(t)-trail:] if trail else ""))
        else:
            out.append(t)
    s = "".join(out)
    s = re.sub(r'\*\*\s*\*\*', '', s)          # מיזוג ריצות סמוכות
    s = re.sub(r'\*\*(\s+)\*\*', r'\1', s)
    return s.strip()

def esc(s):
    return s.replace('\\', '\\\\').replace('"', '\\"')

# תיקוני הקלדה שנכנסו בעריכה הידנית בוורד. הם מוחלים בייבוא ולא ב-content.js,
# כדי שלא יימחקו בייבוא הבא. אם תתקנו אותם בוורד עצמו — השורות כאן פשוט לא יתפסו.
FIXUPS = [
 ("**איך להשפיע על התנהגות בארגוןf**", "**איך להשפיע על התנהגות בארגון**"),
 ("האצלת סמכויות**▪**", "האצלת סמכויות"),
 ("**גישה נורמטיבית המציגה ארגון,▪ תפיסת הארגון**", "**תפיסת הארגון** — גישה נורמטיבית"),
]

def fixup(s):
    for a, b in FIXUPS:
        s = s.replace(a, b)
    return s

doc = docx.Document(SRC)
secs, cur, blocks, items = [], None, None, None
used_sections = set()

def flush_items():
    global items
    if items:
        blocks.append(("items", items)); items = None

def open_sec(name):
    global cur, blocks, items
    flush_items()
    if cur: secs.append((cur, blocks))
    color, sub = SEC_BY_NAME.get(name, ("#37474f", ""))
    cur, blocks, items = (name, color, sub), [], None

for ch in doc.element.body.iterchildren():
    tag = ch.tag.split('}')[1]
    if tag == 'tbl':
        tb = Table(ch, doc)
        rows = [[rich(c.paragraphs[0]) if c.paragraphs else "" for c in r.cells] for r in tb.rows]
        rows = [r for r in rows if any(x.strip() for x in r)]
        if rows and cur:
            flush_items()
            head = rows[0] if all(x.strip() for x in rows[0][:1]) is not None else None
            blocks.append(("tbl", (rows[0], rows[1:])))
        continue
    if tag != 'p': continue
    p = Paragraph(ch, doc)

    # תמונות: דיאגרמות וצילומי טבלאות מהמצגות. הן יושבות בתוך פסקה,
    # ולעיתים בפסקה שכל תוכנה הוא התמונה — ולכן נבדקות לפני בדיקת הטקסט.
    blips = ch.findall('.//' + BLIP)
    if blips and cur is not None:
        for b in blips:
            rid = b.get(EMBED)
            part = doc.part.rels[rid].target_part if rid in doc.part.rels else None
            if part is None: continue
            try:
                uri, w, h = encode_image(part.blob)
            except Exception as e:
                print("!! תמונה נכשלה (%s): %s" % (rid, type(e).__name__)); continue
            flush_items(); blocks.append(("img", (uri, w, h)))

    txt = fixup(rich(p))
    if not txt: continue
    plain = re.sub(r'\*\*', '', txt).strip()
    r0 = p.runs[0] if p.runs else None
    sz = r0.font.size.pt if (r0 and r0.font.size) else None

    # שם מקטע פותח מקטע רק בפעם הראשונה. בסוף החוברת שמות המקטעים חוזרים
    # ככותרות-משנה שמקבצות את שאלות המבחן לפי נושא ("אדיג'ס", "בס ואבוליו"),
    # ובלעדי התנאי הזה כל אחת מהן הייתה פותחת מקטע חדש ועמוד חדש.
    name = ALIAS.get(plain, plain)
    if name in SEC_BY_NAME and name not in used_sections:
        used_sections.add(name); open_sec(name); continue
    if cur is None:                              # שער / כותרת-על לפני המקטע הראשון
        continue
    if plain.startswith('▪'):                    # פריט
        # ה-▪ נכתב לפעמים בתוך ריצה מודגשת, ואז rich() מחזיר "**▪ מונח**".
        # בלי הניקוי הזה ה-** וה-▪ נדבקים לראש הערך וזולגים למפתח הא״ב.
        rt = txt.strip()
        rt = re.sub(r'^\*\*\s*▪\s*', '**', rt)   # "**▪ x**"  ->  "**x**"
        rt = re.sub(r'^▪\s*', '', rt)            # "▪ x"      ->  "x"
        rt = re.sub(r'^\*\*\s*\*\*', '', rt).strip()
        parts = rt.split(' — ', 1)
        if len(parts) == 2:
            entry = "%s :: %s" % (parts[0].strip(), parts[1].strip())
        else:
            entry = rt
        if items is None: items = []
        items.append(entry); continue
    if sz == 10.5 or (sz is None and len(plain) < 46 and not plain.endswith('.')):
        flush_items(); blocks.append(("h2", txt)); continue
    # פסקת המשך — נשמרת כטקסט חופשי
    flush_items(); blocks.append(("p", txt))

flush_items()
if cur: secs.append((cur, blocks))

# ---------- כתיבה ----------
L = ['/* תוכן חוברת החומר הפתוח — מנהיגות בניהול.',
     '   נבנה אוטומטית מקובץ הוורד על ידי from_docx.py — אין לערוך ידנית:',
     '   ערכו את הוורד והריצו מחדש, אחרת העריכה תידרס בייבוא הבא.',
     '',
     '   טקסט עשיר:  **מודגש**  __הדגשה צהובה__  «ציטוט»  ⟨מספר⟩',
     '   ⟦כינוי⟧ בכותרת פריט — לא מוצג, אך נכנס למפתח הא״ב.            */',
     '{sections:[', '']
for (name, color, sub), blocks in secs:
    L.append('/* ══════════════ %s */' % name)
    L.append('{name:"%s", color:"%s", sub:"%s", cols:2, blocks:[' % (esc(name), color, esc(sub)))
    for kind, val in blocks:
        if kind == 'h2':
            L.append('{h2:"%s"},' % esc(val))
        elif kind == 'p':
            L.append('{p:"%s"},' % esc(val))
        elif kind == 'img':
            uri, w, h = val
            L.append('{img:"%s",w:%d,h:%d},' % (uri, w, h))
        elif kind == 'tbl':
            head, rows = val
            L.append('{tbl:{head:[%s],rows:[' % ",".join('"%s"' % esc(c) for c in head))
            for r in rows:
                L.append('  [%s],' % ",".join('"%s"' % esc(c) for c in r))
            L.append(']}},')
        else:
            L.append('{items:[')
            for it in val:
                L.append('"%s",' % esc(it))
            L.append(']},')
    L.append(']},'); L.append('')
L.append(']}')
io.open(OUT, "w", encoding="utf-8").write("\n".join(L))

n_items = sum(len(v) for _, b in secs for k, v in b if k == 'items')
n_tbl = sum(1 for _, b in secs for k, v in b if k == 'tbl')
n_img = sum(1 for _, b in secs for k, v in b if k == 'img')
print("נכתב:", OUT)
print("מקטעים: %d · פריטים: %d · טבלאות: %d · תמונות: %d" % (len(secs), n_items, n_tbl, n_img))
for (n, _, _), b in secs:
    print("   %-22s פריטים %3d" % (n, sum(len(v) for k, v in b if k == 'items')))
