# -*- coding: utf-8 -*-
"""ייצוא חוברת החומר הפתוח לקובץ Word לעריכה.

ה-HTML הוא התוצר להדפסה; ה-Word נועד למי שרוצה לשנות ניסוח, להוסיף
או להוריד — ולהדפיס מוורד. שתי עמודות ושוליים צרים, כדי שגם אחרי
עריכה קלה הקובץ יישאר בסביבות 5 דפים (10 עמודים).

הרצה:  node -e "...content.json..."  &&  python to_docx.py
"""
import io, json, os, re, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import docx
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
COURSE = os.path.dirname(REPO)
DATA = json.load(io.open(os.path.join(HERE, "content.json"), encoding="utf-8"))
OUT = [os.path.join(REPO, "חומר פתוח - מנהיגות בניהול.docx"),
       os.path.join(COURSE, "חומר פתוח - מנהיגות בניהול.docx")]

TOKEN = re.compile(r'(\*\*.+?\*\*|__.+?__|«.+?»|⟨.+?⟩)')

def rtl(par):
    """Word אינו מסיק כיווניות מהטקסט — צריך להצהיר עליה בכל פסקה."""
    pPr = par._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rich(par, text, base=9.5, color=None):
    """מפרק את סימוני הטקסט העשיר של content.js לריצות Word."""
    text = re.sub(r'\s*⟦[^⟧]*⟧', '', text)          # כינויי מפתח — לא מוצגים
    for tok in TOKEN.split(text):
        if not tok: continue
        b = u = False; val = tok
        if tok.startswith('**'):   val, b = tok[2:-2], True
        elif tok.startswith('__'): val, b, u = tok[2:-2], True, True
        elif tok.startswith('«'):  val = '"' + tok[1:-1] + '"'
        elif tok.startswith('⟨'):  val, b = tok[1:-1], True
        r = par.add_run(val)
        r.bold = b; r.underline = u
        r.font.size = Pt(base)
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        if color: r.font.color.rgb = color

def two_columns(section, n=2):
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(n)); cols.set(qn('w:space'), '340')

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Arial'; st.font.size = Pt(9.5)
st.element.rPr.rFonts.set(qn('w:cs'), 'Arial')
st.paragraph_format.space_after = Pt(2)
st.paragraph_format.line_spacing = 1.0

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, a, Mm(11))
sec.right_to_left = True

# ---------- שער ----------
h = doc.add_paragraph(); rtl(h)
r = h.add_run("חומר פתוח — מנהיגות בניהול")
r.bold = True; r.font.size = Pt(17); r.font.name = 'Arial'
r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
sub = doc.add_paragraph(); rtl(sub)
add_rich(sub, "נבנה מ«סיכום קורס מנהיגות מקוצר למבחן». "
              "נושא המוטיבציה ירד מהמבחן ואינו נכלל. "
              "**קובץ זה נועד לעריכה** — הגרסה להדפסה היא קובץ ה-HTML.", 9)
two_columns(sec, 2)

for s in DATA["sections"]:
    ph = doc.add_paragraph(); rtl(ph)
    ph.paragraph_format.space_before = Pt(8)
    r = ph.add_run(s["name"])
    r.bold = True; r.font.size = Pt(12.5); r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    r.font.color.rgb = RGBColor.from_string(s.get("color", "#333333").lstrip('#').upper())

    for b in s["blocks"]:
        if "h2" in b or "h3" in b:
            p = doc.add_paragraph(); rtl(p)
            p.paragraph_format.space_before = Pt(5)
            r = p.add_run(b.get("h2") or b.get("h3"))
            r.bold = True; r.font.size = Pt(10.5); r.font.name = 'Arial'
            r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
            r.font.color.rgb = RGBColor.from_string(s.get("color", "#333333").lstrip('#').upper())
        elif "warn" in b:
            p = doc.add_paragraph(); rtl(p)
            add_rich(p, "⚠ " + b["warn"], 9, RGBColor(0xA5, 0x1B, 0x1B))
        elif "tbl" in b:
            t = b["tbl"]; rows = t.get("rows", []); head = t.get("head")
            if not rows: continue
            tab = doc.add_table(rows=0, cols=len(rows[0]))
            tab.style = 'Table Grid'
            tab.table_direction = 1                   # RTL: העמודה הראשונה מימין
            if head:
                cells = tab.add_row().cells
                for i, c in enumerate(head[:len(rows[0])]):
                    cells[i].paragraphs[0].clear() if hasattr(cells[i].paragraphs[0], 'clear') else None
                    rtl(cells[i].paragraphs[0]); add_rich(cells[i].paragraphs[0], c, 8.5)
                    for rr in cells[i].paragraphs[0].runs: rr.bold = True
            for row in rows:
                cells = tab.add_row().cells
                for i, c in enumerate(row):
                    rtl(cells[i].paragraphs[0]); add_rich(cells[i].paragraphs[0], c, 8.5)
        elif "items" in b:
            for it in b["items"]:
                p = doc.add_paragraph(); rtl(p)
                p.paragraph_format.space_after = Pt(1.5)
                parts = str(it).split(" :: ")
                head = re.sub(r'\s*⟦[^⟧]*⟧', '', parts[0])
                r = p.add_run("▪ " + head)
                r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Arial'
                r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
                if len(parts) > 1:
                    d = p.add_run(" — ")
                    d.font.size = Pt(9.5); d.font.name = 'Arial'
                    d._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
                    add_rich(p, " :: ".join(parts[1:]), 9.5)

for path in OUT:
    doc.save(path)
    print("נכתב:", path)
